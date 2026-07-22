import os
import time
from typing import Optional
from .base import TextExtractor, ProgressCallback
from ..domain.document import (
    Document, DocumentSource, DocumentContent, ExtractionInfo,
    ExtractionMethod, Paragraph, Sentence
)
from ..text_preprocessing import clean_text, extract_paragraphs, tokenize_sentences
from ..utils import FileProcessingError

try:
    import docx
except ImportError:
    docx = None

class DocxExtractor(TextExtractor):
    @property
    def supported_extensions(self) -> tuple[str, ...]:
        return ('.docx',)

    def can_handle(self, file_path: str) -> bool:
        return file_path.lower().endswith(self.supported_extensions)

    def extract(
        self,
        file_path: str,
        progress_callback: Optional[ProgressCallback] = None,
    ) -> Document:
        if docx is None:
            raise FileProcessingError("python-docx is not installed. Cannot read DOCX.")

        start_time = time.time()
        
        try:
            doc = docx.Document(file_path)
            raw_text = "\n".join([p.text for p in doc.paragraphs])
        except Exception as e:
            raise FileProcessingError(f"Error reading DOCX {file_path}: {e}")

        clean_txt = clean_text(raw_text)
        para_texts = extract_paragraphs(clean_txt)
        
        paragraphs = []
        total_words = 0
        for i, pt in enumerate(para_texts):
            sents = tokenize_sentences(pt)
            sentences = []
            char_offset = 0
            for j, s in enumerate(sents):
                idx = pt.find(s, char_offset)
                if idx == -1: idx = char_offset
                sentences.append(Sentence(s, j, idx, idx + len(s)))
                char_offset = idx + len(s)
            
            p_words = len(pt.split())
            total_words += p_words
            paragraphs.append(Paragraph(
                text=pt,
                index=i,
                page_number=0,
                word_count=p_words,
                char_count=len(pt),
                is_ocr_derived=False,
                sentences=tuple(sentences)
            ))

        content = DocumentContent(
            raw_text=raw_text,
            paragraphs=tuple(paragraphs),
            word_count=total_words,
            paragraph_count=len(paragraphs),
            sentence_count=sum(p.sentence_count for p in paragraphs)
        )

        source = DocumentSource(
            file_path=file_path,
            file_name=os.path.basename(file_path),
            extension=".docx",
            file_size_bytes=os.path.getsize(file_path),
            mtime=os.path.getmtime(file_path)
        )

        extraction_info = ExtractionInfo(
            method=ExtractionMethod.DOCX,
            page_count=0,
            ocr_page_count=0,
            extraction_time_s=time.time() - start_time,
            warnings=()
        )

        return Document(source=source, content=content, extraction_info=extraction_info)
