"""
test_one_to_many.py

Backend tests for the One-to-Many cross-checking orchestrator.

Covers:
    1. One-to-One comparison still works (regression, using the same engine)
    2. One-to-Many with 2 documents
    3. One-to-Many with 5+ documents
    4. A document with high similarity
    5. A document with low/no similarity
    6. One invalid document while other documents remain valid
    7. Empty document list
    8. Duplicate documents
    9. OCR-based (scanned) documents
    10. Real end-to-end One-to-Many with the genuine embedding engine

The orchestration tests use a fast, deterministic stub engine registered into
the same EngineRegistry so they do not depend on downloading the AI model.
"""

import os
import tempfile

from backend.assignment_analyzer import AssignmentAnalyzer
from backend.engines import EngineRegistry, EngineConfig, ComparisonEngine, EngineCapabilities
from backend.engines.base import ProgressCallback
from backend.domain.document import Document
from backend.domain.comparison import (
    ComparisonResult, SimilarityStatistics, SimilarityBand,
    MatchedParagraph,
)

STUB_ID = "stub_v1"

REFERENCE_PARAS = [
    "Photosynthesis is the process by which plants convert sunlight into chemical energy.",
    "DNA carries the genetic instructions used in the growth and functioning of all living organisms.",
]

HIGH_SIM_PARAS = [
    "Photosynthesis is the process by which plants convert sunlight into chemical energy.",
    "DNA carries the genetic instructions used in the growth and functioning of all living organisms.",
]

PARTIAL_PARAS = [
    "Photosynthesis converts sunlight into chemical energy within plants.",
    "DNA encodes genetic instructions for growth and functioning of organisms.",
]

LOW_SIM_PARAS = [
    "The stock market fluctuates with quarterly earnings reports and interest rates.",
    "Ancient civilizations developed complex irrigation systems along river valleys.",
]

NO_SIM_PARAS = [
    "Quantum tunneling is a purely quantum mechanical phenomenon.",
    "Cheese ripening is accelerated by specific bacterial cultures.",
]


class StubEngine(ComparisonEngine):
    """Deterministic, dependency-free engine with the same output contract as
    the embedding engine: paragraphs matched by token-overlap similarity,
    aggregated into coverage * average score."""

    ENGINE_ID = STUB_ID
    ENGINE_NAME = "Stub Deterministic Engine"

    @property
    def engine_id(self) -> str:
        return self.ENGINE_ID

    @property
    def engine_name(self) -> str:
        return self.ENGINE_NAME

    def is_available(self) -> bool:
        return True

    @property
    def capabilities(self) -> EngineCapabilities:
        return EngineCapabilities(
            paragraph_matching=True,
            sentence_matching=False,
            span_matching=False,
            multilingual=False,
            offline=True,
            requires_gpu=False,
            approximate_time_per_para=0.0001,
        )

    @staticmethod
    def _tokens(text: str) -> set:
        import re
        return set(re.findall(r"[a-z0-9']+", text.lower()))

    def compare(
        self,
        doc_a: Document,
        doc_b: Document,
        config: EngineConfig,
        progress_callback: ProgressCallback = None,
    ) -> ComparisonResult:
        matched = []
        matched_b = set()
        for pa in doc_a.paragraphs:
            wa = self._tokens(pa.text)
            best_score = 0.0
            best_pb = None
            for i, pb in enumerate(doc_b.paragraphs):
                if i in matched_b:
                    continue
                wb = self._tokens(pb.text)
                inter = len(wa & wb)
                union = len(wa | wb)
                score = inter / union if union else 0.0
                if score > best_score:
                    best_score, best_pb = score, (i, pb)
            if best_score >= config.similarity_threshold and best_pb is not None:
                matched_b.add(best_pb[0])
                matched.append((best_score, pa, best_pb[1]))

        matched_paras = [
            MatchedParagraph(pa, pb, score, ())
            for score, pa, pb in matched
        ]

        matched_idx_a = {p.index for _, p, _ in matched}
        matched_idx_b = {q.index for _, _, q in matched}
        total_paras = len(doc_a.paragraphs) + len(doc_b.paragraphs)
        coverage = (len(matched_idx_a) + len(matched_idx_b)) / total_paras if total_paras else 0.0
        avg = sum(s for s, _, _ in matched) / len(matched) if matched else 0.0
        overall = coverage * avg

        stats = SimilarityStatistics(
            overall_score=overall,
            score_percent=min(100, int(overall * 100)),
            max_match_score=max((s for s, _, _ in matched), default=0.0),
            min_match_score=min((s for s, _, _ in matched), default=0.0),
            avg_match_score=avg,
            coverage_a=len(matched_idx_a) / len(doc_a.paragraphs) if doc_a.paragraphs else 0.0,
            coverage_b=len(matched_idx_b) / len(doc_b.paragraphs) if doc_b.paragraphs else 0.0,
            band=SimilarityBand.from_score(overall),
            confidence=0.95,
        )

        return ComparisonResult(
            mode="one_to_one",
            engine_id=self.ENGINE_ID,
            doc_a=doc_a,
            doc_b=doc_b,
            statistics=stats,
            matched_paragraphs=tuple(matched_paras),
            unique_paragraphs_a=tuple(
                p for p in doc_a.paragraphs if p.index not in matched_idx_a
            ),
            unique_paragraphs_b=tuple(
                p for p in doc_b.paragraphs if p.index not in matched_idx_b
            ),
            metadata_warnings=(),
            summary="",
            processing_time_s=0.001,
            error=False,
            error_message=None,
        )


# Register once, exactly like the built-in engines, so the orchestrator can
# resolve it through the same EngineRegistry used in production.
EngineRegistry.register(StubEngine)


def _docx(path: str, paragraphs: list[str]):
    from docx import Document as DocxDocument
    doc = DocxDocument()
    for p in paragraphs:
        doc.add_paragraph(p)
    doc.save(path)


def _scanned_pdf(path: str, page_count: int = 1):
    import fitz
    pdf = fitz.open()
    for _ in range(page_count):
        page = pdf.new_page(width=595, height=842)
        pix = fitz.Pixmap(fitz.csGRAY, fitz.IRect(0, 0, 220, 220))
        pix.clear_with(255)
        page.insert_image(fitz.Rect(50, 50, 300, 300), pixmap=pix)
    pdf.save(path)
    pdf.close()


def _analyzer(config: EngineConfig = None) -> AssignmentAnalyzer:
    return AssignmentAnalyzer(
        config=config or EngineConfig(
            similarity_threshold=0.4,
            enable_sentence_matching=False,
        ),
        engine_id=STUB_ID,
    )


# ---------------------------------------------------------------------------
# 1. One-to-One still works
# ---------------------------------------------------------------------------
def test_one_to_one_regression():
    with tempfile.TemporaryDirectory() as tmp:
        a = os.path.join(tmp, "a.docx")
        b = os.path.join(tmp, "b.docx")
        _docx(a, HIGH_SIM_PARAS)
        _docx(b, PARTIAL_PARAS)
        result = _analyzer().analyze_one_to_one(a, b)
        assert not result.get("error"), result.get("summary")
        assert result["score"] > 0
        print("test_one_to_one_regression: score =", result["score"])


# ---------------------------------------------------------------------------
# 2. One-to-Many with 2 documents
# ---------------------------------------------------------------------------
def test_one_to_many_two_docs():
    with tempfile.TemporaryDirectory() as tmp:
        ref = os.path.join(tmp, "ref.docx")
        s1 = os.path.join(tmp, "s1.docx")
        s2 = os.path.join(tmp, "s2.docx")
        _docx(ref, REFERENCE_PARAS)
        _docx(s1, HIGH_SIM_PARAS)
        _docx(s2, NO_SIM_PARAS)
        result = _analyzer().analyze_one_to_many(ref, [s1, s2])
        assert not result.get("error"), result.get("summary")
        pairs = result["pairs"]
        assert len(pairs) == 2
        assert pairs[0]["score"] > pairs[1]["score"]
        assert result["score"] == pairs[0]["score"]  # highest
        assert result["similar_paragraphs"] > 0
        print("test_one_to_many_two_docs:", [p["score"] for p in pairs])


# ---------------------------------------------------------------------------
# 3. One-to-Many with 5+ documents  (deterministic ordering)
# ---------------------------------------------------------------------------
def test_one_to_many_five_docs():
    with tempfile.TemporaryDirectory() as tmp:
        ref = os.path.join(tmp, "ref.docx")
        _docx(ref, REFERENCE_PARAS)
        targets = [
            HIGH_SIM_PARAS, PARTIAL_PARAS, LOW_SIM_PARAS, NO_SIM_PARAS, HIGH_SIM_PARAS,
        ]
        paths = []
        for i, paras in enumerate(targets):
            p = os.path.join(tmp, f"s{i + 1}.docx")
            _docx(p, paras)
            paths.append(p)
        result = _analyzer().analyze_one_to_many(ref, paths)
        assert not result.get("error"), result.get("summary")
        pairs = result["pairs"]
        assert len(pairs) == 5
        # Deterministic order == input order
        assert [p["error"] for p in pairs] == [False] * 5
        assert all(not p["error"] for p in pairs)
        scores = [p["score"] for p in pairs]
        assert scores[0] > scores[1] > scores[3]
        assert result["score"] == max(scores)
        print("test_one_to_many_five_docs:", scores)


# ---------------------------------------------------------------------------
# 4/5. High similarity and low/no similarity
# ---------------------------------------------------------------------------
def test_high_similarity():
    with tempfile.TemporaryDirectory() as tmp:
        ref = os.path.join(tmp, "ref.docx")
        high = os.path.join(tmp, "high.docx")
        _docx(ref, REFERENCE_PARAS)
        _docx(high, HIGH_SIM_PARAS)
        result = _analyzer().analyze_one_to_many(ref, [high])
        assert not result.get("error")
        assert result["pairs"][0]["score"] == 100
        print("test_high_similarity: 100? ->", result["pairs"][0]["score"])


def test_no_similarity():
    with tempfile.TemporaryDirectory() as tmp:
        ref = os.path.join(tmp, "ref.docx")
        none = os.path.join(tmp, "none.docx")
        _docx(ref, REFERENCE_PARAS)
        _docx(none, NO_SIM_PARAS)
        result = _analyzer().analyze_one_to_many(ref, [none])
        assert not result.get("error")
        assert result["pairs"][0]["score"] == 0
        assert result["score"] == 0
        print("test_no_similarity: 0? ->", result["pairs"][0]["score"])


# ---------------------------------------------------------------------------
# 6. One invalid document while others remain valid
# ---------------------------------------------------------------------------
def test_invalid_document_isolated():
    with tempfile.TemporaryDirectory() as tmp:
        ref = os.path.join(tmp, "ref.docx")
        s1 = os.path.join(tmp, "s1.docx")
        s2 = os.path.join(tmp, "s2.docx")
        missing = os.path.join(tmp, "does_not_exist.docx")
        _docx(ref, REFERENCE_PARAS)
        _docx(s1, HIGH_SIM_PARAS)
        _docx(s2, PARTIAL_PARAS)
        result = _analyzer().analyze_one_to_many(ref, [s1, missing, s2])
        pairs = result["pairs"]
        assert len(pairs) == 3
        # Individual failure does NOT crash the batch or fail it overall
        assert not result.get("error"), result.get("summary")
        assert pairs[1]["error"] is True
        assert pairs[1]["score"] == 0
        assert pairs[0]["score"] > 0 and pairs[2]["score"] > 0
        assert result["score"] > 0
        assert "could not be completed" in result.get("summary", "")
        print("test_invalid_document_isolated:", [p["error"] for p in pairs])


# ---------------------------------------------------------------------------
# 7. Empty document list
# ---------------------------------------------------------------------------
def test_empty_document_list():
    with tempfile.TemporaryDirectory() as tmp:
        ref = os.path.join(tmp, "ref.docx")
        _docx(ref, REFERENCE_PARAS)
        result = _analyzer().analyze_one_to_many(ref, [])
        assert result.get("error") is True
        assert result["score"] == 0
        assert "No comparison documents" in result.get("summary", "")
        print("test_empty_document_list: handled gracefully")


# ---------------------------------------------------------------------------
# 8. Duplicate documents (deterministic, both included, in order)
# ---------------------------------------------------------------------------
def test_duplicate_documents():
    with tempfile.TemporaryDirectory() as tmp:
        ref = os.path.join(tmp, "ref.docx")
        dupe = os.path.join(tmp, "dupe.docx")
        _docx(ref, REFERENCE_PARAS)
        _docx(dupe, HIGH_SIM_PARAS)
        result = _analyzer().analyze_one_to_many(ref, [dupe, dupe])
        pairs = result["pairs"]
        assert len(pairs) == 2
        assert pairs[0]["score"] == pairs[1]["score"]
        assert pairs[0]["score"] == 100
        print("test_duplicate_documents:", [p["score"] for p in pairs])


# ---------------------------------------------------------------------------
# 9. OCR-based (scanned) document isolation
# ---------------------------------------------------------------------------
def test_scanned_doc_does_not_crash_batch():
    with tempfile.TemporaryDirectory() as tmp:
        ref = os.path.join(tmp, "ref.docx")
        s1 = os.path.join(tmp, "s1.docx")
        scanned = os.path.join(tmp, "scanned.pdf")
        _docx(ref, REFERENCE_PARAS)
        _docx(s1, HIGH_SIM_PARAS)
        _scanned_pdf(scanned)
        result = _analyzer().analyze_one_to_many(ref, [s1, scanned])
        pairs = result["pairs"]
        assert len(pairs) == 2
        assert pairs[0]["score"] > 0  # normal doc compares fine
        assert not result.get("error")  # batch overall not failed
        print("test_scanned_doc_does_not_crash_batch: scanned pair error? ->",
              pairs[1]["error"])


# ---------------------------------------------------------------------------
# 10. Real end-to-end with the genuine embedding engine
# ---------------------------------------------------------------------------
def test_real_embedding_one_to_many():
    try:
        EngineRegistry.get("embedding_v1")
    except Exception as e:
        print(f"test_real_embedding_one_to_many: SKIPPED (embedding engine "
              f"unavailable: {e})")
        return

    with tempfile.TemporaryDirectory() as tmp:
        ref = os.path.join(tmp, "ref.docx")
        high = os.path.join(tmp, "high.docx")
        zero = os.path.join(tmp, "zero.docx")
        _docx(ref, REFERENCE_PARAS)
        _docx(high, HIGH_SIM_PARAS)
        _docx(zero, NO_SIM_PARAS)

        analyzer = AssignmentAnalyzer(
            config=EngineConfig(enable_sentence_matching=True),
            engine_id="embedding_v1",
        )
        result = analyzer.analyze_one_to_many(ref, [high, zero])
        assert not result.get("error"), result.get("summary")
        pairs = result["pairs"]
        assert len(pairs) == 2
        assert pairs[0]["score"] > pairs[1]["score"]
        assert not any(p["error"] for p in pairs)
        print("test_real_embedding_one_to_many:", [p["score"] for p in pairs])


if __name__ == "__main__":
    tests = [
        ("test_one_to_one_regression", test_one_to_one_regression),
        ("test_one_to_many_two_docs", test_one_to_many_two_docs),
        ("test_one_to_many_five_docs", test_one_to_many_five_docs),
        ("test_high_similarity", test_high_similarity),
        ("test_no_similarity", test_no_similarity),
        ("test_invalid_document_isolated", test_invalid_document_isolated),
        ("test_empty_document_list", test_empty_document_list),
        ("test_duplicate_documents", test_duplicate_documents),
        ("test_scanned_doc_does_not_crash_batch", test_scanned_doc_does_not_crash_batch),
        ("test_real_embedding_one_to_many", test_real_embedding_one_to_many),
    ]
    for name, fn in tests:
        print(f"\n--- {name} ---")
        try:
            fn()
        except AssertionError as e:
            print(f"FAIL: {e}")
            raise
    print("\nAll one-to-many tests passed!")